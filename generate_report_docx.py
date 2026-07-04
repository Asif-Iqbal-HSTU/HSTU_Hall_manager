import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import os

md_path = r"c:\inetpub\wwwroot\hall_manager_hstu\HSTU_Hall_Manager_Report.md"
docx_path = r"c:\inetpub\wwwroot\hall_manager_hstu\HSTU_Hall_Manager_Report.docx"

doc = Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Helper to add headings with nice colors and fonts
def add_custom_heading(text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 78, 121) # Dark Indigo/Blue
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(46, 117, 182) # Medium Blue
        run.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(89, 89, 89) # Dark Grey
        run.bold = True
    return heading

# Helper to format inline markdown style
def format_run(p, text):
    clean_text = text
    # Clean up double stars or single stars
    parts = re.split(r'(\*\*.*?\*\*)', clean_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            bold_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', bold_text)
            r = p.add_run(bold_text)
            r.font.name = 'Segoe UI'
            r.bold = True
        else:
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    italic_text = subpart[1:-1]
                    italic_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', italic_text)
                    r = p.add_run(italic_text)
                    r.font.name = 'Segoe UI'
                    r.italic = True
                else:
                    subpart = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', subpart)
                    subpart = subpart.replace('$', '') # clean up math signs
                    r = p.add_run(subpart)
                    r.font.name = 'Segoe UI'

# Parse MD file lines
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_headers = []
table_rows = []
in_code = False
code_lines = []

i = 0
while i < len(lines):
    line = lines[i].strip('\n')
    trimmed = line.strip()
    
    # Code block check
    if trimmed.startswith('```'):
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("\n".join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(128, 0, 0)
            in_code = False
            code_lines = []
        else:
            in_code = True
        i += 1
        continue
        
    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Table parsing
    if trimmed.startswith('|') and not in_table:
        in_table = True
        table_headers = [c.strip() for c in trimmed.split('|')[1:-1]]
        # skip separator line
        i += 2
        continue
    
    if in_table:
        if trimmed.startswith('|'):
            row_cols = [c.strip() for c in trimmed.split('|')[1:-1]]
            table_rows.append(row_cols)
            i += 1
            continue
        else:
            # Table ended, construct it
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Light Shading Accent 1'
            
            hdr_cells = table.rows[0].cells
            for col_idx, text in enumerate(table_headers):
                hdr_cells[col_idx].text = text
                for p in hdr_cells[col_idx].paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for run in p.runs:
                        run.font.bold = True
                        run.font.name = 'Segoe UI'
            
            for row_data in table_rows:
                row = table.add_row()
                for col_idx, text in enumerate(row_data):
                    if col_idx < len(row.cells):
                        row.cells[col_idx].text = text
                        for p in row.cells[col_idx].paragraphs:
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.space_before = Pt(2)
                            for run in p.runs:
                                run.font.name = 'Segoe UI'
            
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_after = Pt(6)
            
            in_table = False
            table_headers = []
            table_rows = []
            # Fallthrough to process the current line if it contains content
    
    if in_table:
        i += 1
        continue

    # Title & Headers
    if trimmed.startswith('# '):
        # Main Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(trimmed[2:])
        run.font.name = 'Segoe UI'
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.bold = True
    elif trimmed.startswith('## '):
        add_custom_heading(trimmed[3:], 1, space_before=18, space_after=8)
    elif trimmed.startswith('### '):
        add_custom_heading(trimmed[4:], 2, space_before=14, space_after=6)
    elif trimmed.startswith('#### '):
        add_custom_heading(trimmed[5:], 3, space_before=10, space_after=4)
    # Bullet points
    elif trimmed.startswith('- ') or trimmed.startswith('* '):
        text = trimmed[2:]
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        format_run(p, text)
    # Numbered points
    elif re.match(r'^\d+\.\s', trimmed):
        m = re.match(r'^(\d+\.)\s(.*)', trimmed)
        num_prefix = m.group(1)
        text = m.group(2)
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        format_run(p, text)
    # Normal Paragraph
    elif trimmed:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        format_run(p, trimmed)
        
    i += 1

doc.save(docx_path)
print("Docx created successfully at:", docx_path)
